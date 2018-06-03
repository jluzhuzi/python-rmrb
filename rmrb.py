
#这是我的第一个程序，现学现卖。目的是抓取人民网上的评论员文章。
#人民网上有四个列表页，这四页是固定不变的，提供了所有可抓取的网址。
#流程为：
#1.先列出四个列表页；
#2.从四个列表中抽取文章的URL；
#3.从文章网页提取标题和正文；
#4.写入空白Word文件。
##############
######使用方法：######
#先安装python3；
#再安装pip，用pip安装bs4、python-docx这两个库。
#在终端运行python3 RenMinRiBao.py
#在本文件同级目录下打开wenji.docx即可。
#wenji.docx排版会不太好看，这是因为样式处理起来很麻烦，我偷懒了，就是这么理直气壮。
#我的微信：zhuzicn，一名自由讲师，主讲申论、公基、结构化面试，哲学硕士，不是学计算机的，所以谢绝类似“代码太Lower了”这样的评价。。。。
##############
#先导入库文件
from urllib.request import urlopen
from bs4 import BeautifulSoup
import docx #这个库安装的时候名字叫python-docx，但使用时变成这个名字。
wenji = docx.Document()#生成一个空白的Docx文件
wangzhiliebiao=[]#生成一个空白列表，用来存储文章网址
#下面四个是四个列表页，可以穷尽所有文章
laiyuanwangye01=urlopen("http://opinion.people.com.cn/GB/8213/353915/354347/index1.html")
laiyuanwangye02=urlopen("http://opinion.people.com.cn/GB/8213/353915/354347/index2.html")
laiyuanwangye03=urlopen("http://opinion.people.com.cn/GB/8213/353915/354347/index3.html")
laiyuanwangye04=urlopen("http://opinion.people.com.cn/GB/8213/353915/354347/index4.html")
#用一个很可笑的方式对四个列表页进行操作，得到所有文章的网址。第一次写，不知道如何写得更简洁。这样程序也能跑，不是吗？
laiyuanwangzhi01=BeautifulSoup(laiyuanwangye01,"lxml")
for wangzhi in laiyuanwangzhi01.find("td",{"class":"t11"}).findAll("a"):
    if "href" in wangzhi.attrs:
        wangzhiliebiao.append(wangzhi.attrs['href'])#第一个列表页里的网址被装进列表了。
laiyuanwangzhi02=BeautifulSoup(laiyuanwangye02,"lxml")
for wangzhi in laiyuanwangzhi02.find("td",{"class":"t11"}).findAll("a"):
    if "href" in wangzhi.attrs:
        wangzhiliebiao.append(wangzhi.attrs['href'])#第二个列表页里的网址被装进列表了。

laiyuanwangzhi03=BeautifulSoup(laiyuanwangye03,"lxml")
for wangzhi in laiyuanwangzhi03.find("td",{"class":"t11"}).findAll("a"):
    if "href" in wangzhi.attrs:
        wangzhiliebiao.append(wangzhi.attrs['href'])#第三个列表页里的网址被装进列表了。
laiyuanwangzhi04=BeautifulSoup(laiyuanwangye04,"lxml")
for wangzhi in laiyuanwangzhi04.find("td",{"class":"t11"}).findAll("a"):
    if "href" in wangzhi.attrs:
        wangzhiliebiao.append(wangzhi.attrs['href'])#第四个列表页里的网址被装进列表了。
for dangewangzhi in wangzhiliebiao:
    html=urlopen("http://opinion.people.com.cn"+dangewangzhi) #原网址为相对网址，现在补全，成为绝对地址
    wangye = BeautifulSoup(html.read(),"lxml")
    biaoti = wangye.findAll({"h1"})
    for dabiaoti in biaoti:
        wenji.add_heading(dabiaoti.get_text(),1)#遍历写入标题
    wenzhang = wangye.findAll("div",{"class":"box_con"})
    for wenzhangtext in wenzhang:
        wenji.add_paragraph(wenzhangtext.get_text())#遍历写入正文
wenji.save("./wenji.docx")#保存文件，完工
